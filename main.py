print("TURTLES' RALLY LIEPĀJA")
#vārdi, sākums

print("Spēles izstrādātāji: Paula Lībeka, Loreta Aleksejeva, Keita Laimiņa")
print("SPIED 'ENTER', LAI SĀKTU SPĒLI")

#taustiņu pārbaude
input() 
print("starts")
#import keyboard
#while True:
    #if keyboard.is_pressed("space"):
        #print("STARTS")
        #break
# Reading an excel file using Python
#2dmasīvs ar jautājumiem
jautajumi = {
  'a':{ 'jautajums':"Kurš ir koncertzāles'Lielais Dzintars' arhitekts?\na:Folkers Gīnke \nb:Gunars Birkerts\nc:Pauls Makss Berči", 'pareiza_atbilde': ['a']},
  'b':{ 'jautajums':"Kādā krāsā ir Liepājas muzejs?\na:zils \nb:rozā\nc:zaļš", 'pareiza_atbilde': ['c']},
  'cf':{'jautajums':"Kā sauc šo zivi?\na:plekste\nb:menca\nc:līdaka", 'pareiza_atbilde': ['a'], 'jautajums2':"Kā sauc šo zivi?\na:asaris\nb:lasis\nc:rauda", 'pareiza_atbilde2':['b']},
  'd':{'jautajums':"Kā vārdos ir nosaukti Liepājas tirgi?\na:Katrīnas un Jāņa \nb:Annas un Pētera\nc:Marijas un Jāņa", 'pareiza_atbilde':['b']},
  'ef':{'jautajums':"Kā sauc šo prezidentu(-i)?", 'pareiza_atbilde':["vaira vīķe freiberga","vaira vīķe-freiberga","vaira vīķe - freiberga","vaira vīķe -freiberga", "vaira vīķe -freiberga"], 'jautajums2':"Kā sauc šo prezidentu(-i)?", 'pareiza_atbilde2':['gustavs zemgals'], 'pareiza_atbilde3':['alberts kviesis']},
  'g':{'jautajums':"Kādas ticības baznīca atrodas Karostā?\na:katoļu \nb:luterāņu\nc:pareizticīgo", 'pareiza_atbilde':['c']},
  'h':{'jautajums':"Kura ir Grobiņas lielākā iela?\na:Brīvības iela \nb:Lielā iela\nc:Saules iela", 'pareiza_atbilde':['b']},
  'i':{'jautajums':"Kādus dzīvniekus var novērot zirgu salā?\na:putnus \nb:zirgus\nc:roņus", 'pareiza_atbilde':['a']},
  'j':{'jautajums':"Kādā secībā jāveic pirmā palīdzība?(atbilžu burtus atdali ar komatu)\na:sauc palīgā \nb:pārliecinies par savu un apkārtējo drošību.\nc:pārbaudi cietušā stāvokli", 'pareiza_atbilde':['b,c,a']}
}
print(jautajumi['a']['jautajums'])

#laikapstākļu ģenerators
import random
import string
apstakli=[2, 1]
laikapstakli = random.choice(apstakli)
if laikapstakli ==1:
  print("Laikapstākļi ir labi")
if laikapstakli == 2:
  print("Laikapstākļi ir slikti")

global punkti
punkti = 10 

#kontrolpunktu faila nolasīšana
with open('faili/kontrolpunkti.txt') as f:
    contents = f.read()
    print(contents)

#import docx
#doc = docx.Document('faili/Kontrolpunkts_muzejs.docx')
#paras = [p.text for p in doc.paragraphs if p.text] 
 
#word faila nolasīšana
import docx
doci = docx.Document('faili/Kontrolpunkts_muzejs.docx')
#len(doci.paragraphs)
print(doci.paragraphs[0-3].text)
#print(doci.paragraphs[1].text)
#print(doci.paragraphs[2].text)
#print(doci.paragraphs[3].text)
#len(doci.paragraphs[1].runs)
#print(doci.paragraphs[0].runs[0].text)
#open excel
import pandas as pd
df = pd.read_excel (r'faili/kontrolpunktu tabula (bez km).xlsx')
data=pd.read_excel (r'faili/kontrolpunktu tabula (bez km).xlsx')
dklk = pd.DataFrame(data, columns= ['a','b'])
print (df)
import xlrd
 
# Give the location of the file
loc = ("faili/kontrolpunktu tabula (bez km).xlsx")
 
# To open Workbook
wb = xlrd.open_workbook(loc)
sheet = wb.sheet_by_index(0)
print(sheet.cell_value(1,14))
#vietas indeksi
indeksi = {
  'a':1,'b':2,'c':3,'d':4, 'e':5, 'f':6, 'g':7,'h':8,'i':9,'j':10,'1': 11, '2':12, '3':13, 'starts':14, 'finišs':14
  }
#vietas nosaukumi
nosaukumi = {
  'a':'Lielais Dzintars','b':'Liepājas muzejs','c':'Kursa','d':'Pētertirgus', 'e':'Čakstes laukums', 'f':'Juliannas pagalms', 'g':'Karostas katedrāle','h':'Grobiņas pilskalns','i':'Zirgu sala','j':'Liepājas slimnīca','1': 'uzlādes stacija "Virši"', '2':'uzlādes stacija "Circle K"', '3':'uzlādes stacija "Viada"', 'starts':'Liepājas Olimpiskais centrs', 'finišs':'Ziemeļu mols'
}
#kontrolpunkts Kursa
def kontrolpunktsC():
  print("Kā sauc šo zivi?\na:plekste\nb:menca\nc:līdaka")
  from PIL import Image
  im = Image.open(r"zivis/plekste.jpg")
  im.show()
  zivs=input()
  if zivs=="a":
      global punkti
      punkti= punkti+10
      print("Pareizi!")
  else:
      print("Mēģini vēlreiz!")
      print("Kā sauc šo zivi?\na:asaris\nb:lasis\nc:rauda")
      from PIL import Image
      im = Image.open(r"zivis/lasis.jpg")
      im.show() 
      zivs=input()
      if zivs=="b":
          punkti=punkti+5
          print("Pareizi!")
      else:
          print("Nepareizi!")
#kontrolpunkts cakstes laukumus
def kontrolpunktsE():
  print("Kā sauc šo prezidentu(-i)?") 
  from PIL import Image
  im = Image.open(r"prezidenti/Zemgals.jpg")
  im.show()
  atbilde =input()
  if atbilde == "gustavs zemgals":
      global punkti
      punkti=punkti+10
      print("Pareizi!")
  else:
      print("Mēģini vēlreiz!")
      from PIL import Image 
      im = Image.open(r"prezidenti/Janis_Cakste.jpg")
      im.show()
      print("Kā sauc šo prezidentu(-i)?") 
      atbilde=input()
      if atbilde== "jānis čakste":
          punkti=punkti+5
          print("Pareizi!")
      else:
          print("Mēģini vēlreiz!")
          from PIL import Image
          im = Image.open(r"prezidenti/kviesis.png")
          im.show()
          print("Kā sauc šo prezidentu(-i)?") 
          atbilde=input()
          if atbilde=="Alberts Kviesis"or atbilde=="alberts kviesis":
              print("Pareizi!")#VAJAG BONUSA PUNKTUS VĒL?
              punkti+=3
          else:
            from PIL import Image
            im = Image.open(r"prezidenti/vaira.jpg")
            im.show()
            atbilde =input()
            pareizi=["Vaira Vīķe Freiberga","Vaira Vīķe-Freiberga","Vaira Vīķe - Freiberga","vaira vīķe freiberga","vaira vīķe-freiberga", "vaira vīķe - freiberga"]
            if atbilde in pareizi:
                print("Pareizi!")
                punkti+=2
            else:
                print("Nepareizi!")
#kontrolpunkts Juliannas pagalms f
def kontrolpunktsF():
  global punkti
  for i in range(5): 
    gads=random.randrange(1990,2031)
    menesis=random.randrange(1,13)
    if menesis in range(1,10):
        menesis2="0"+str(menesis)
    if menesis in range(10,13):
        menesis2=menesis
    if menesis==2:
        diena=random.randrange(1,29)
    elif menesis==4 or menesis==6 or menesis==9 or menesis==11:
        diena=random.randrange(1,31)
    else:
        diena=random.randrange(1,32)
    if diena in range(1,10):
        diena1="0"+str(diena)
    if diena in range(10,32):
        diena1=diena

    gads2=random.randrange(1980, gads-10)
    menesis3=random.randrange(1,menesis+1)
    if menesis3 in range(1,10):
        menesis4="0"+str(menesis3)
    if menesis3 in range(10,13):
        menesis4=menesis3
    else:
        menesis=menesis3 
    if menesis3==2:
        diena2=random.randrange(1,diena+1)
    elif menesis3==4 or menesis3==6 or menesis3==9 or menesis3==11:
        diena2=random.randrange(1,diena+1)
    else:
        diena2=random.randrange(1,diena+1)
    if diena2 in range(1,10):
        diena="0"+str(diena2)
    else:
        diena=diena2
    print("Tu esi bāra apsargs. Šodien ir šāds datums: "+str(diena1)+"."+str(menesis2)+"."+str(gads))
    print("Vai tu drīksti bārā ielaist cilvēku, kura dzimšanas datums: "+str(diena)+"."+str(menesis4)+"."+str(gads2))
    print("a:drīkst ielaist\nb:nedrīkst ielaist")
    apsargs = input()
    if gads - gads2>18:
      if apsargs=="a":
        punkti=punkti+10
        print("Pareizi!")
      else:
        print("Nepareizi!")
    if gads - gads2<18:
      if apsargs=="b":
        punkti=punkti+10
        print("Pareizi!")
      else:
        print("Nepareizi!")
    if int(gads) - int(gads2)==18 and int(diena1) - int(diena)>0 and int(menesis2) - int(menesis4)>=0 :
      if apsargs=="a":
        punkti=punkti+10
        print("Pareizi!")
      else:
        print("Nepareizi!")

#akumulatora izlāde
akumulators = 100
def izlade_akumulators(laikapstakli,cels):
  #uztaisīt, lai reālus km var mainīt
  if laikapstakli == 1 and int(cels)<=140:
    akumulators_izteretais = cels*100/140 
    print(akumulators_izteretais)
  if laikapstakli==1 and int(cels)>140:
    akumulators_izteretais=akumulators
    print("Akumulators izlādējies")
  if laikapstakli == 2 and int(cels)<=80:
    akumulators_izteretais= cels*100/80 
    print(akumulators_izteretais)
  if laikapstakli==2 and int(cels)>80:
    akumulators_izteretais= akumulators
    print("Akumulators izlādējies")
  return akumulators_izteretais
#evakuators
def evakuators():
  global vieta_tagad
  global punkti
  vieta_tagad = input("Akumulators ir izlādējies! Nepieciešams evakuators! Par evakuatora izmantošanu tiks noņemti 20 bonusa punkti. Ievadi uzlādes stacijas ciparu(1-3): ")
  punkti = punkti-20 
  print("Tu atrodies: "+nosaukumi[vieta_tagad])

#uzlādes stacijas
def uzlāde():
  global vieta_tagad
  global punkti
  global akumulators
  print("Tu atrodies:", nosaukumi[vieta_tagad],"\nVienu bonusa punktu var izmantot, lai uzlādētu vienu procentu akumulatora.")
  akumulators_uzladets = akumulators
  l=0
  while akumulators_uzladets == akumulators:
    print("Tev ir sakrāti:", str(punkti),"punkti. Cik bonusa punktus vēlies izmantot uzlādei?")
    punkti_uzladei = input()
    if punkti_uzladei == 'pārbaude':
      cels=0
      akumulators = round(akumulators-izlade_akumulators(laikapstakli, cels))
      kilometri = 140*akumulators/100
      print("uzlādes līmenis ir "+str(akumulators)+"%\nvēl var nobraukt "+str(kilometri)+" km")
    #print("Tu atrodies:", nosaukumi[vieta_tagad],"\nTev ir sakrāti:", str(punkti),"punkti.\nVienu bonusa punktu var izmantot, lai uzlādētu vienu procentu akumulatora.\nCik bonusa punktus vēlies izmantot?")
    #punkti_uzladei = input()
    elif int(punkti_uzladei) > 0 and int(punkti_uzladei) <= punkti:
      #akumulators += int(punkti_uzladei)
      akumulators_uzladets += int(punkti_uzladei)
      punkti = punkti - int(punkti_uzladei)
      print("Akumulatora līmenis pēc uzlādes ir:",str(akumulators_uzladets),"%")
    elif int(punkti_uzladei) < 0 or int(punkti_uzladei) > int(punkti):
      print("Tev ir sakrāti", str(punkti),"!\nIevadi, cik punktus vēlies iztērēt priekš uzlādes: ")
    elif int(punkti_uzladei) == 0:
      akumulators_uzladets=akumulators+1 
      l+=1
  if l!=0:
    akumulators=akumulators_uzladets-1
  if l == 0:
    akumulators=akumulators_uzladets
  print(str(akumulators),"%")
    
    
      
  

#spēles jēga
akumulators = 100
kontrolpunkti = 0
vieta_tagad="starts" 
cels=0
vieta_biji=[]


while kontrolpunkti < 10: 
  #jāuzlabo, lai zina, kur atrodies
  vieta = input("Uz kurieni dosies? Ievadi atbistošo kontrolpunkta burtu: ")
  while vieta in vieta_biji or vieta==vieta_tagad:
    print("Tu jau vietā -",nosaukumi[vieta],"- esi bijis.")
    vieta=input("Uz kurieni dosies? Ievadi atbistošo kontrolpunkta burtu: ")
  #aprēķina akumulatora izlādi atkarībā no laikapstākļiem
  if vieta == 'vieta':
    print(nosaukumi[vieta_tagad])
  elif vieta == 'pārbaude':
    cels=0
    akumulators = round(akumulators-izlade_akumulators(laikapstakli, cels))
    kilometri = 140*akumulators/100
    print("uzlādes līmenis ir "+str(akumulators)+"%\nvēl var nobraukt "+str(kilometri)+" km")
  elif vieta in jautajumi:
    vieta_biji.append(vieta)
    cels = sheet.cell_value(indeksi[vieta],indeksi[vieta_tagad])
    print(cels)
    kontrolpunkti = kontrolpunkti+1
    akumulators = round(akumulators-izlade_akumulators(laikapstakli, cels))
    kilometri = 140*akumulators/100
    if akumulators <=0:
      evakuators()
    #if atbilde in jautajumi[vieta]['pareiza_atbilde']:
     # punkti=punkti+10
      #print("Pareizi!")
    if akumulators>0:
      print(jautajumi[vieta]['jautajums'])
      atbilde = input()
      k=0
      while atbilde not in jautajumi[vieta]['pareiza_atbilde']:
        k+=1
        print("Nepareizi!Mēģini vēlreiz!")
        print(jautajumi[vieta]['jautajums'])
        atbilde = input()
      print("Pareizi!")
      punkti+=10/(k+1) if k<2 else 0
      vieta_tagad =vieta
  elif vieta == 'f':
    kontrolpunktsF()
    cels = sheet.cell_value(indeksi[vieta],indeksi[vieta_tagad])
    print(cels)
    vieta_biji.append(vieta)
    vieta_tagad =vieta
    kontrolpunkti+=1
  elif vieta == 'e':
    kontrolpunktsE()
    cels = sheet.cell_value(indeksi[vieta],indeksi[vieta_tagad])
    print(cels)
    vieta_biji.append(vieta)
    vieta_tagad =vieta
    kontrolpunkti+=1
  elif vieta == 'c':
    kontrolpunktsC()
    cels = sheet.cell_value(indeksi[vieta],indeksi[vieta_tagad])
    print(cels)
    vieta_biji.append(vieta)
    vieta_tagad =vieta
    kontrolpunkti+=1
  elif vieta == '1' or vieta=='2' or vieta == '3':
    vieta_tagad = vieta
    uzlāde()
  else:
    print("Tu ievadīji neesošu kontrolpunktu, uzlādes staciju vai komandu.")
  

  #if input("ievadi a")=="a":
    #kontrolpunkti= kontrolpunkti+1
#while True:
    #if keyboard.is_pressed("a"):
        #print("Uzlādes līmenis ir: "+str(akumulators)+"%")
        #break
#indentifikatora atbilstības pārbaude
print("Tev ir bonusa punkti:",str(punkti))


#akumolatora uzskaite, a0=100%, vienmēr var uzspiest


#laikapstākļu faktors mašīnas nobraukumam


#uzlādes līmeņa, km nolasīšana jebkurā laikā


#10 kotrolpunkti, uzlādes stacija, uzdevumi


#bonusa punkti auto uzlādei par atrisinātu uzdevumu


#evakuators
if akumulators == 0:
  print("Akumulators ir izlādējies! Nepieciešams evakuators! Ievadi uzlādes stacijas identifikatoru!")
  uzlades_stacija = input()
  punkti = punkti-20
#spēles beigas, rezultātu izvade
#pavadītais laiks spēlē
#bildes, uzlabot, lai randoma atveras un pazūd
#kontrolpunkts Juliannas pagalms PABEIGT?padomāt kā var vairākas reizes bez tik daudz koda(loop)

print(punkti)
    


